import psycopg2
import pandas as pd
import requests
import json
import uuid
from datetime import datetime, timedelta, date
import docx
import tempfile
import os
import subprocess
from typing import List, Optional, Dict
from loguru import logger

# Глобальная переменная для хранения time_inserted по claim_id
time_inserted_cache = {}


def quick_db_query():
    """Быстрый запрос к БД"""

    # Параметры подключения (замените на свои)
    conn_string = "host=10.10.4.172 port=5432 dbname=frap user=postgres password=VYgJBx6Eun}W"

    try:
        # Подключаемся и выполняем запрос
        conn = psycopg2.connect(conn_string)

        # Получаем данные в DataFrame
        df = pd.read_sql_query(
            "SELECT * FROM frap_smev_declaration_cert_checks where declaration_certificate_id is null", conn)

        # Выводим информацию
        logger.debug(f"Всего записей: {len(df)}")
        logger.debug("\nПервые 10 записей:")
        logger.debug(df.head(10))

        logger.debug(df.head())

        return df

    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return None


def get_frap_info(df: pd.DataFrame) -> List[int]:
    """
    Обрабатывает DataFrame с записями frap_smev_declaration_cert_checks,
    проверяет статусы в таблице frap_claims и отправляет запросы для изменения статуса.

    Args:
        df: DataFrame с данными из frap_smev_declaration_cert_checks

    Returns:
        List[int]: Список frap_claims.id с статусом 2002 после обработки
    """

    # Параметры подключения к БД (должны быть такие же как в quick_db_query)
    conn_string = "host=10.10.4.172 port=5432 dbname=frap user=postgres password=VYgJBx6Eun}W"

    # URL для получения токена и изменения статуса
    token_url = "http://lk-test.test-kuber-nd.fsrar.ru/api-lc-frap/tools/token?role=developer"
    status_url_template = "http://lk-test.test-kuber-nd.fsrar.ru/api-lc-frap/dashboard/frap-claims/{}/status"

    # Сначала соберем все claims_id, которые будем обрабатывать
    claims_to_process = []

    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()

        # Получаем уникальные frap_claims_id из DataFrame
        frap_claims_ids = df['frap_claims_id'].dropna().unique().tolist()

        logger.info(f"Найдено {len(frap_claims_ids)} уникальных frap_claims_id для обработки")

        for frap_claims_id in frap_claims_ids:
            try:
                # Ищем запись в frap_claims с максимальным id и проверяем дату
                query = """
                SELECT id, status_id, date_inserted, time_inserted 
                FROM frap_claims 
                WHERE id = %s 
                AND date_inserted >= CURRENT_DATE - INTERVAL '1 days'
                ORDER BY id DESC 
                LIMIT 1
                """

                cursor.execute(query, (frap_claims_id,))
                result = cursor.fetchone()

                if result:
                    claim_id, status_id, date_inserted, time_inserted = result

                    if status_id == 2001:
                        logger.info(f"Статус 2001 найден для claim_id: {claim_id}, отправляем запрос...")
                        # Сохраняем time_inserted в кэш
                        time_inserted_cache[claim_id] = time_inserted
                        claims_to_process.append(claim_id)
                    else:
                        logger.debug(f"Статус не 2001 для claim_id: {claim_id}, пропускаем")
            except Exception as e:
                logger.error(f"Ошибка при проверке claim_id {frap_claims_id}: {e}")
                continue

    except Exception as e:
        logger.error(f"Ошибка подключения к БД: {e}")
        return []
    finally:
        if 'conn' in locals():
            conn.close()

    # Отправляем запросы для всех claims со статусом 2001
    for claim_id in claims_to_process:
        try:
            # Получаем bearer токен
            token_response = requests.get(token_url)

            if token_response.status_code == 200:
                bearer_token = token_response.text.strip()
                logger.debug("Токен успешно получен")

                # Формируем URL для изменения статуса
                status_url = status_url_template.format(claim_id)

                # Заголовки для запроса
                headers = {
                    "accept": "*/*",
                    "Authorization": f"{bearer_token}",
                    "Content-Type": "application/json"
                }

                # Данные для отправки
                data = {
                    "comment": "",
                    "statusCode": [2002]
                }

                # Отправляем PATCH запрос
                patch_response = requests.patch(
                    status_url,
                    headers=headers,
                    data=json.dumps(data)
                )

                logger.debug(f"Ответ API для claim_id {claim_id}: {patch_response.status_code} - {patch_response.text}")

        except Exception as e:
            logger.error(f"Ошибка при отправке запроса для claim_id {claim_id}: {e}")

    # Теперь проверяем, какие claims действительно получили статус 2002
    processed_claims = []
    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()

        if claims_to_process:
            check_query = """
            SELECT id FROM frap_claims 
            WHERE id IN ({}) AND status_id = 2002
            and date_inserted >= CURRENT_DATE - INTERVAL '1 days'
            """.format(','.join(['%s'] * len(claims_to_process)))

            cursor.execute(check_query, claims_to_process)
            results = cursor.fetchall()

            processed_claims = [result[0] for result in results]
            logger.info(f"Фактически получили статус 2002: {len(processed_claims)} claims")
        else:
            logger.info("Нет claims для проверки статуса 2002")

    except Exception as e:
        logger.error(f"Ошибка при проверке финального статуса: {e}")
    finally:
        if 'conn' in locals():
            conn.close()

    return processed_claims


def process_declaration_certificates(processed_claims_ids: List[int]):
    """
    Функция для обработки declaration_certificate для конкретных frap_claims.id

    Args:
        processed_claims_ids: Список frap_claims.id, для которых был установлен статус 2002
    """
    conn_string = "host=10.10.4.172 port=5432 dbname=frap user=postgres password=VYgJBx6Eun}W"

    try:
        conn = psycopg2.connect(conn_string)

        if not processed_claims_ids or len(processed_claims_ids) == 0:
            logger.info("Нет обработанных claims для второго этапа")
            return None

        logger.info(f"Работаем с {len(processed_claims_ids)} claims, для которых установлен статус 2002")

        # Находим записи в declaration_certificate с type_doc = DECLARATION для обработанных claims
        query_declaration = """
        SELECT id, frap_claims_id, number, date_validity
        FROM declaration_certificate 
        WHERE type_doc = 'DECLARATION'
        AND frap_claims_id IN ({})
        """.format(','.join(['%s'] * len(processed_claims_ids)))

        df_declaration = pd.read_sql_query(query_declaration, conn, params=processed_claims_ids)
        logger.info(f"Найдено {len(df_declaration)} записей в declaration_certificate для обработанных claims")

        if len(df_declaration) == 0:
            logger.info("Нет подходящих записей в declaration_certificate для обработанных claims")
            return None

        # Выборка А: id, frap_claims_id (явно преобразуем типы)
        selection_a = df_declaration[['id', 'frap_claims_id']].copy()
        selection_a['id'] = selection_a['id'].astype(int)
        selection_a['frap_claims_id'] = selection_a['frap_claims_id'].astype(int)

        logger.success(f"Успешно получено {len(selection_a)} записей для обработки")
        return selection_a

    except Exception as e:
        logger.error(f"Ошибка при обработке declaration_certificates: {e}")
        return None
    finally:
        if 'conn' in locals():
            conn.close()


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Конвертирует DOCX в PDF используя LibreOffice
    """
    try:
        # Пробуем использовать LibreOffice для конвертации
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path), docx_path
        ], capture_output=True, text=True, timeout=30)

        if result.returncode == 0:
            # LibreOffice создает файл с тем же именем но расширением .pdf
            expected_pdf = docx_path.replace('.docx', '.pdf')
            if os.path.exists(expected_pdf):
                if expected_pdf != pdf_path:
                    os.rename(expected_pdf, pdf_path)
                return True
        return False
    except Exception as e:
        logger.error(f"Ошибка при конвертации через LibreOffice: {e}")
        return False


def generate_smev_pdf(declaration_certificate_id):
    """
    Генерирует PDF файл на основе docx шаблона и возвращает hex представление и даты

    Args:
        declaration_certificate_id: ID записи в declaration_certificate

    Returns:
        tuple: (hex_content, request_datetime, response_datetime) или None в случае ошибки
    """
    conn_string = "host=10.10.4.172 port=5432 dbname=frap user=postgres password=VYgJBx6Eun}W"

    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()

        # Получаем данные из declaration_certificate и связанных таблиц
        query = """
        SELECT 
            dc.id as decl_id,
            dc.number as decl_number,
            dc.date_validity,
            fc.id as frap_claim_id,
            fc.doc_number,
            fc.doc_date,
            fc.reg_number,
            fc.reg_date,
            fc.reg_time,
            fs.uuid as smev_uuid,
            fs.datetime_created
        FROM declaration_certificate dc
        JOIN frap_claims fc ON dc.frap_claims_id = fc.id
        LEFT JOIN frap_smev_declaration_cert_checks fs ON fs.declaration_certificate_id = dc.id
        WHERE dc.id = %s
        """

        cursor.execute(query, (declaration_certificate_id,))
        result = cursor.fetchone()

        if not result:
            logger.error(f"Запись не найдена для declaration_certificate_id: {declaration_certificate_id}")
            return None

        # Распаковываем результат
        (decl_id, decl_number, date_validity, frap_claim_id,
         doc_number, doc_date, reg_number, reg_date, reg_time,
         smev_uuid, datetime_created) = result

        # Загружаем и редактируем docx файл
        doc = docx.Document('2105.docx')

        # Функция для замены текста в документе
        def replace_text_in_doc(doc, old_text, new_text):
            # Преобразуем new_text в строку, если это не строка
            if not isinstance(new_text, str):
                new_text = str(new_text) if new_text is not None else ""

            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)

        # Форматируем даты
        def format_date(date_value):
            if date_value:
                if isinstance(date_value, str):
                    try:
                        date_obj = datetime.strptime(date_value, '%Y-%m-%d')
                        return date_obj.strftime('%d.%m.%Y')
                    except:
                        return date_value
                elif isinstance(date_value, (datetime, date)):
                    return date_value.strftime('%d.%m.%Y')
            return date_value if date_value else ""

        # Форматируем datetime
        def format_datetime(dt_value):
            if dt_value:
                if isinstance(dt_value, str):
                    try:
                        dt_obj = datetime.strptime(dt_value, '%Y-%m-%d %H:%M:%S')
                        return dt_obj.strftime('%d.%m.%Y, %H:%M:%S')
                    except:
                        return dt_value
                elif isinstance(dt_value, datetime):
                    return dt_value.strftime('%d.%m.%Y, %H:%M:%S')
            return dt_value if dt_value else ""

        # Функция для форматирования времени (убираем микросекунды)
        def format_time(time_value):
            if time_value:
                if isinstance(time_value, str):
                    # Если время содержит микросекунды, обрезаем до секунд
                    if '.' in time_value:
                        time_value = time_value.split('.')[0]
                    # Если время содержит миллисекунды, обрезаем до секунд
                    elif ',' in time_value:
                        time_value = time_value.split(',')[0]
                    return time_value
                elif isinstance(time_value, datetime):
                    return time_value.strftime('%H:%M:%S')
            return time_value if time_value else ""

        # Создаем даты для запроса и ответа
        if datetime_created:
            request_datetime = datetime_created
            response_datetime = datetime_created + timedelta(seconds=30)
        else:
            # Если datetime_created нет, используем текущее время
            current_time = datetime.now()
            request_datetime = current_time
            response_datetime = current_time + timedelta(seconds=30)

        # Заменяем данные в документе
        replace_text_in_doc(doc, '314-734-234', str(doc_number) if doc_number else "")
        replace_text_in_doc(doc, '13.11.2025', format_date(doc_date))
        replace_text_in_doc(doc, '03-00005787', str(reg_number) if reg_number else "")

        # Формируем полную строку для даты и времени фиксации в ЕГАИС
        # Используем сохраненное time_inserted из кэша
        if frap_claim_id in time_inserted_cache:
            saved_time_inserted = time_inserted_cache[frap_claim_id]
            if saved_time_inserted:
                # Если time_inserted - это время (без даты), комбинируем с doc_date
                if isinstance(saved_time_inserted, str) and len(saved_time_inserted.split(':')) >= 2:
                    # time_inserted содержит только время, форматируем его и комбинируем с датой из doc_date
                    formatted_time = format_time(saved_time_inserted)
                    if doc_date:
                        if isinstance(doc_date, (datetime, date)):
                            date_part = doc_date.strftime('%d.%m.%Y')
                        else:
                            date_part = format_date(doc_date)
                        egais_full_datetime = f"{date_part}, {formatted_time}"
                    else:
                        egais_full_datetime = f"{formatted_time}"
                else:
                    # time_inserted содержит полную дату-время
                    if isinstance(saved_time_inserted, (datetime, date)):
                        egais_date = saved_time_inserted.strftime('%d.%m.%Y')
                        egais_time = saved_time_inserted.strftime('%H:%M:%S')  # Уже без микросекунд
                        egais_full_datetime = f"{egais_date}, {egais_time}"
                    else:
                        egais_full_datetime = str(saved_time_inserted)
            else:
                egais_full_datetime = ""
        else:
            egais_full_datetime = ""

        logger.info(f"Используем time_inserted для frap_claim_id {frap_claim_id}: {egais_full_datetime}")
        replace_text_in_doc(doc, '13.11.2025, 16:15:14', egais_full_datetime)

        # Заменяем данные в таблицах
        replace_text_in_doc(doc, '5082891d-e736-41d0-a1e6-1c9df6a531fa',
                            str(smev_uuid) if smev_uuid else str(uuid.uuid4()))

        # Используем наши даты для PDF - заменяем ВСЕ вхождения дат и времен
        replace_text_in_doc(doc, '14.11.2025, 16:27:02', format_datetime(request_datetime))
        replace_text_in_doc(doc, '14.11.2025, 13:17:02', format_datetime(request_datetime))

        replace_text_in_doc(doc, 'ЕАЭС N RU Д-FR.АГ82.В.07662/27', str(decl_number) if decl_number else "")
        replace_text_in_doc(doc, '11.03.2022', format_date(date_validity))

        # Сведения об ответе
        new_response_uuid = str(uuid.uuid4())
        replace_text_in_doc(doc, 'f0b8b65b-41ea-42ca-adb4-2effd5a64029', new_response_uuid)

        replace_text_in_doc(doc, '14.11.2025, 16:28:56', format_datetime(response_datetime))

        replace_text_in_doc(doc, 'не имеет', 'имеет')

        # ДОПОЛНИТЕЛЬНО: Заменяем отдельно время если оно встречается отдельно
        # Используем сохраненное time_inserted
        if frap_claim_id in time_inserted_cache:
            saved_time_inserted = time_inserted_cache[frap_claim_id]
            if saved_time_inserted:
                if isinstance(saved_time_inserted, str) and len(saved_time_inserted.split(':')) >= 2:
                    # time_inserted содержит только время, форматируем его
                    egais_time_only = format_time(saved_time_inserted)
                elif isinstance(saved_time_inserted, (datetime, date)):
                    # time_inserted содержит полную дату-время, извлекаем только время
                    egais_time_only = saved_time_inserted.strftime('%H:%M:%S')
                else:
                    egais_time_only = format_time(str(saved_time_inserted))

                replace_text_in_doc(doc, '16:15:14', egais_time_only)

        # Также заменяем возможные другие форматы времени
        replace_text_in_doc(doc, '16:27:02', request_datetime.strftime('%H:%M:%S'))
        replace_text_in_doc(doc, '13:17:02', request_datetime.strftime('%H:%M:%S'))
        replace_text_in_doc(doc, '16:28:56', response_datetime.strftime('%H:%M:%S'))

        # Сохраняем временный docx файл
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
            temp_doc_path = temp_doc.name
            doc.save(temp_doc_path)

        # Конвертируем в PDF используя LibreOffice
        temp_pdf_path = temp_doc_path.replace('.docx', '.pdf')

        if convert_docx_to_pdf(temp_doc_path, temp_pdf_path):
            logger.info("Успешная конвертация через LibreOffice")
        else:
            # Если LibreOffice не установлен, создаем простой текстовый PDF
            logger.warning("LibreOffice не доступен, создаем простой PDF")
            create_simple_pdf(temp_pdf_path, doc)

        # Читаем PDF и конвертируем в hex
        with open(temp_pdf_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()
            hex_content = pdf_content.hex()

        # Удаляем временные файлы
        os.unlink(temp_doc_path)
        if os.path.exists(temp_pdf_path):
            os.unlink(temp_pdf_path)

        logger.success(f"PDF успешно сгенерирован для declaration_certificate_id: {declaration_certificate_id}")
        return hex_content, request_datetime, response_datetime

    except Exception as e:
        logger.error(f"Ошибка при генерации PDF для declaration_certificate_id {declaration_certificate_id}: {e}")
        return None
    finally:
        if 'conn' in locals():
            conn.close()


def create_simple_pdf(pdf_path, doc):
    """
    Создает простой PDF файл если LibreOffice не доступен
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter

        # Простой текст для PDF
        text = "Документ сгенерирован автоматически\n"
        text += "Оригинальный DOCX файл был изменен но не может быть конвертирован в PDF\n"
        text += "Установите LibreOffice для полной функциональности"

        c.drawString(100, height - 100, text)
        c.save()

    except ImportError:
        # Если reportlab не установлен, создаем пустой файл
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4\n')


def save_smev_check_file(declaration_certificate_id, hex_content, request_datetime, response_datetime):
    """
    Сохраняет файл проверки в таблицу frap_smev_check_files
    и обновляет запись в frap_smev_declaration_cert_checks в два этапа

    Args:
        declaration_certificate_id: ID записи в declaration_certificate
        hex_content: hex представление PDF файла
        request_datetime: datetime для поля datetime_created (из PDF "Дата и время запроса")
        response_datetime: datetime для поля datetime_delivered (из PDF "Дата и время ответа")
    """
    conn_string = "host=10.10.4.172 port=5432 dbname=frap user=postgres password=VYgJBx6Eun}W"

    try:
        conn = psycopg2.connect(conn_string)
        cursor = conn.cursor()

        # Явно преобразуем в int
        decl_id = int(declaration_certificate_id)

        # Получаем frap_claims_id из declaration_certificate
        get_frap_claim_query = """
        SELECT frap_claims_id FROM declaration_certificate WHERE id = %s
        """
        cursor.execute(get_frap_claim_query, (decl_id,))
        frap_claim_result = cursor.fetchone()

        if not frap_claim_result:
            logger.error(f"Не найден frap_claims_id для declaration_certificate_id: {decl_id}")
            return False

        frap_claims_id = frap_claim_result[0]

        # ЭТАП 1: Обновляем datetime_created в frap_smev_declaration_cert_checks
        update_datetime_created_query = """
        UPDATE frap_smev_declaration_cert_checks 
        SET datetime_created = %s
        WHERE frap_claims_id = %s
        """

        # Форматируем дату для БД
        request_datetime_str = request_datetime.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]
        cursor.execute(update_datetime_created_query, (request_datetime_str, frap_claims_id))
        logger.info(f"Обновлен datetime_created для frap_claims_id: {frap_claims_id}")

        # Проверяем, существует ли уже запись в frap_smev_check_files
        check_query = """
        SELECT id FROM frap_smev_check_files 
        WHERE smev_check_id = %s AND smev_check_type_id = 3
        """
        cursor.execute(check_query, (decl_id,))
        existing_record = cursor.fetchone()

        if existing_record:
            logger.warning(f"Запись уже существует для smev_check_id: {decl_id}")
            return False

        # Вставляем новую запись в frap_smev_check_files
        insert_query = """
        INSERT INTO public.frap_smev_check_files 
        (smev_check_id, smev_check_type_id, pdf, frap_claim_cor_id) 
        VALUES(%s, 3, decode(%s, 'hex'), NULL)
        RETURNING id
        """

        cursor.execute(insert_query, (decl_id, hex_content))
        smev_check_file_id = cursor.fetchone()[0]
        logger.info(f"Создана запись в frap_smev_check_files с id: {smev_check_file_id}")

        # ЭТАП 2: Обновляем остальные поля в frap_smev_declaration_cert_checks
        update_final_query = """
        UPDATE frap_smev_declaration_cert_checks 
        SET declaration_certificate_id = %s,
            datetime_delivered = %s,
            is_processed = true,
            result = true
        WHERE frap_claims_id = %s
        """

        # Форматируем дату для БД
        response_datetime_str = response_datetime.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]
        cursor.execute(update_final_query, (decl_id, response_datetime_str, frap_claims_id))

        conn.commit()
        logger.success(f"Запись успешно сохранена для declaration_certificate_id: {decl_id}")
        logger.success(f"Обновлена запись в frap_smev_declaration_cert_checks для frap_claims_id: {frap_claims_id}")
        return True

    except Exception as e:
        logger.error(
            f"Ошибка при сохранении smev_check_file для declaration_certificate_id {declaration_certificate_id}: {e}")
        if 'conn' in locals():
            conn.rollback()
        return False
    finally:
        if 'conn' in locals():
            conn.close()


def unified_processing():
    """
    Объединенная функция, которая последовательно выполняет оба функционала
    """
    logger.info("=== НАЧАЛО ОБЪЕДИНЕННОЙ ОБРАБОТКИ ===")

    # Шаг 1: Исходный функционал - получаем записи и устанавливаем статус 2002
    logger.info("=== ЭТАП 1: Установка статуса 2002 ===")
    df = quick_db_query()
    if df is not None:
        processed_claims = get_frap_info(df)
        logger.info(f"Установлен статус 2002 для claims: {processed_claims}")

        if processed_claims and len(processed_claims) > 0:
            # Шаг 2: Новый функционал - работаем ТОЛЬКО с claims, для которых установлен статус 2002
            logger.info("=== ЭТАП 2: Генерация PDF для claims со статусом 2002 ===")
            selection_a = process_declaration_certificates(processed_claims)

            if selection_a is not None and len(selection_a) > 0:
                logger.info(f"Найдено {len(selection_a)} записей declaration_certificate для claims со статусом 2002")

                # Шаг 3: Для каждой записи генерируем PDF и сохраняем
                for _, row in selection_a.iterrows():
                    decl_id = int(row['id'])
                    frap_claim_id = int(row['frap_claims_id'])

                    logger.info(
                        f"Генерация PDF для declaration_certificate_id: {decl_id} (frap_claim_id: {frap_claim_id})")
                    result = generate_smev_pdf(decl_id)

                    if result:
                        hex_content, request_datetime, response_datetime = result
                        logger.info(f"Сохранение файла для declaration_certificate_id: {decl_id}")
                        save_smev_check_file(decl_id, hex_content, request_datetime, response_datetime)
                    else:
                        logger.error(f"Не удалось сгенерировать PDF для declaration_certificate_id: {decl_id}")
            else:
                logger.info("Нет записей declaration_certificate для claims со статусом 2002")
        else:
            logger.info("Нет claims с установленным статусом 2002 для второго этапа")
    else:
        logger.info("Нет данных для обработки на первом этапе")

    logger.info("=== ОБЪЕДИНЕННАЯ ОБРАБОТКА ЗАВЕРШЕНА ===")


if __name__ == "__main__":
    unified_processing()